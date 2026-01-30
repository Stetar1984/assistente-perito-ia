import io
from datetime import date

import streamlit as st
from PIL import Image, ImageOps
from docx import Document
from docx.shared import Inches

st.set_page_config(page_title="Assistente Perito", layout="wide")
st.title("Assistente Perito")
st.caption("MVP: dati pratica + note + foto con didascalie → DOCX")

# -----------------------------
# Helpers
# -----------------------------
def safe_filename(s: str) -> str:
    s = (s or "").strip()
    cleaned = "".join(c for c in s if c.isalnum() or c in ("-", "_"))
    return cleaned.strip("_-") or "pratica"

def normalize_image(uploaded_file, max_side=1600, quality=85) -> io.BytesIO:
    """
    - Corregge orientamento EXIF
    - Ridimensiona (thumbnail)
    - Converte in JPEG
    """
    img = Image.open(uploaded_file)
    img = ImageOps.exif_transpose(img)  # fix rotazioni smartphone
    if img.mode not in ("RGB",):
        img = img.convert("RGB")
    img.thumbnail((max_side, max_side))

    buf = io.BytesIO()
    img.save(buf, format="JPEG", quality=quality, optimize=True)
    buf.seek(0)
    return buf

def build_docx(n_pratica: str, assicurato: str, data_sopralluogo: str, note: str, photos: list) -> io.BytesIO:
    """
    photos: list di dict {name, caption, img_bytesio}
    """
    doc = Document()
    doc.add_heading(f"Relazione Tecnica - Pratica {n_pratica}", level=0)
    doc.add_paragraph(f"Assicurato: {assicurato or '—'}")
    doc.add_paragraph(f"Data sopralluogo: {data_sopralluogo}")

    doc.add_heading("Descrizione / Note tecniche", level=1)
    doc.add_paragraph(note.strip() if note and note.strip() else "—")

    if photos:
        doc.add_heading("Documentazione fotografica", level=1)
        # Tabella: 2 colonne (foto + didascalia)
        table = doc.add_table(rows=1, cols=2)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Foto"
        hdr_cells[1].text = "Didascalia"

        for p in photos:
            row = table.add_row().cells
            # Inserisci immagine nel primo cell
            paragraph = row[0].paragraphs[0]
            run = paragraph.add_run()
            p["img_bytesio"].seek(0)
            run.add_picture(p["img_bytesio"], width=Inches(2.7))
            # Didascalia
            row[1].text = (p["caption"] or p["name"] or "—")

    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out

# -----------------------------
# UI: dati
# -----------------------------
with st.expander("Dati identificativi", expanded=True):
    c1, c2, c3 = st.columns([1, 1, 1])
    n_pratica = c1.text_input("Numero pratica *")
    assicurato = c2.text_input("Assicurato")
    data_sopralluogo = c3.date_input("Data sopralluogo", value=date.today()).isoformat()

st.subheader("Note tecniche")
note = st.text_area("Trascrizione vocale / appunti", height=140, placeholder="Es: Paraurti anteriore dx deformato...")

st.subheader("Foto danni (con didascalie)")
foto_files = st.file_uploader(
    "Carica foto (JPG/PNG)",
    accept_multiple_files=True,
    type=["jpg", "jpeg", "png"],
)

photo_items = []
if foto_files:
    st.caption(f"Foto caricate: {len(foto_files)}")

    # UI per didascalie + ordine
    for idx, f in enumerate(foto_files, start=1):
        with st.container(border=True):
            colA, colB = st.columns([1, 2])

            # Anteprima
            colA.image(f, width=220, caption=f.name)

            # Metadati
            order = colB.number_input(f"Ordine per {f.name}", min_value=1, max_value=999, value=idx, step=1, key=f"ord_{idx}_{f.name}")
            caption = colB.text_input(f"Didascalia per {f.name}", value="", key=f"cap_{idx}_{f.name}")

            # Prepara immagine normalizzata (lazy: solo se servirà, ma qui la facciamo subito per semplicità)
            img_buf = normalize_image(f)
            photo_items.append({"order": int(order), "name": f.name, "caption": caption, "img_bytesio": img_buf})

    # Ordina per ordine scelto
    photo_items.sort(key=lambda x: x["order"])

st.divider()

# -----------------------------
# Generazione DOCX
# -----------------------------
if st.button("Genera DOCX"):
    if not n_pratica.strip():
        st.error("Numero pratica obbligatorio.")
    else:
        docx_bytes = build_docx(
            n_pratica=n_pratica.strip(),
            assicurato=assicurato.strip(),
            data_sopralluogo=data_sopralluogo,
            note=note,
            photos=photo_items,
        )

        st.download_button(
            label="Scarica relazione .docx",
            data=docx_bytes,
            file_name=f"Perizia_{safe_filename(n_pratica)}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
